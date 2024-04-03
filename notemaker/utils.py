from pytube import YouTube 
import moviepy.editor as mp
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os
import cv2
import numpy as np
from openai import OpenAI
import os

import google.generativeai as genai
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image

import tempfile
import wave

from dotenv import load_dotenv
load_dotenv()


client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
)



PROMPT = '''
Read the text taken from youtube videos audio. 
Organize the content into a structured document with clear headings, bullet points, and subpoints. 
Begin with a title that accurately reflects the central topic of the video. 
Then, for each section of the video, create a heading that encapsulates the main idea. 
Under each heading, use bullet points to summarize the key information, details, arguments, or steps presented.
If there are subtopics within a section, introduce subheadings and list relevant points beneath them. 
Ensure the notes are concise, accurate, and capture the essence of the discussion without verbatim transcription. 
Make sure to include code of programs or some diagram etc if any.
The final document should be well-organized, easy to follow, and serve as a useful study or reference guide
'''

def downloadVideo(link, id):
    yt = YouTube(link)
    yt = yt.streams.get_lowest_resolution()
    try:
        yt.download(output_path='tmp/',filename=f"{id}_{yt.title}.mp4")
    except:
        raise Exception("Error downloading video")

    return f"{id}_{yt.title}.mp4"

def getAudio(filename, id):
    clip = mp.VideoFileClip(f"tmp/{filename}")
    try:
        clip.audio.write_audiofile(f"tmp/{id}_{filename}.wav")
    except:
        raise Exception("Error converting video to audio")
    return f"{id}_{filename}.wav"

def add_list(paragraph, text):
    run = paragraph.add_run(text)
    paragraph.style = 'ListBullet'



def createDocument(txt, filename, two_columns=False):
    doc = Document()

    if two_columns:
        section = doc.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')
    
    lines = txt.split('\n')
    for line in lines:
        if line.startswith('# '):  # Title
            heading = doc.add_heading('', level=0)
            run = heading.add_run(line[2:])
            run.bold = True
        elif line.startswith('## '):  # Subtitle
            heading = doc.add_heading('', level=1)
            run = heading.add_run(line[3:])
            run.bold = True
        elif line.startswith('- '):  # Points
            para = doc.add_paragraph()
            add_list(para, line[2:])
        elif line.startswith('  - '):  # Subpoints
            para = doc.add_paragraph()
            add_list(para, line[4:])
        else:  # Regular text
            doc.add_paragraph(line)
    
    doc.save(f"tmp/{filename}")
    return filename

def calculatePercentageDifference(frame1, frame2):
    # Compute the absolute difference between frames
    diff = cv2.absdiff(frame1, frame2)
    # Calculate the percentage difference
    percentage_diff = (np.sum(diff) / (frame1.size * 255.0)) * 100
    return percentage_diff

def splitVideoFrames(videoPath, outputDir):
    os.makedirs(outputDir, exist_ok=True)

    # Open the video file
    cap = cv2.VideoCapture(f"tmp/{videoPath}")

    # Get frames per second (fps) of the video
    fps = int(cap.get(cv2.CAP_PROP_FPS))

    # Variables to keep track of frames
    frameCount = 0
    lastSavedFrame = None

    while True:
        # Read a frame from the video
        ret, frame = cap.read()

        if not ret:
            break

        # Calculate percentage difference with the last saved frame
        if lastSavedFrame is not None:
            percentage_diff = calculatePercentageDifference(lastSavedFrame, frame)
            # If percentage difference is greater than 0.1%, save the frame
            if percentage_diff > 1:
                frameFilename = os.path.join(outputDir, f'frame_{frameCount:04d}.jpg')
                cv2.imwrite(frameFilename, frame)
                lastSavedFrame = frame
        else:
            # Save the first frame
            frameFilename = os.path.join(outputDir, f'frame_{frameCount:04d}.jpg')
            cv2.imwrite(frameFilename, frame)
            lastSavedFrame = frame

        frameCount += 1
    cap.release()

def imageToText(framePath):
    IMG_PROMPT = ''' 
    Read the text from the image and summarize the content in a few bullet points.
    '''
    res = ""
    genai.configure(api_key=os.getenv("GENAI_API_KEY"))
    for imagePath in os.listdir(framePath):
        img = Image.open(f"{framePath}/{imagePath}")
        model = genai.GenerativeModel("gemini-pro-vision")
        response = model.generate_content([IMG_PROMPT, img])
        res += response.text
        os.remove(f"{framePath}/{imagePath}")
    return res

def generateNotes(res):
    prompt = f"{PROMPT}\n\n{res}"
    completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        model="gpt-3.5-turbo",
    )

    response = completion.choices[0].message.content
    return response

def audioToText(audio_file, max_chunk_size=20 * 1024 * 1024):  # 20 MB in bytes
 

    with open(f"tmp/{audio_file}", "rb") as audio_file_obj:
        file_size = os.path.getsize(f"tmp/{audio_file}")
        num_chunks = (file_size + max_chunk_size - 1) // max_chunk_size
        transcripts = []

        for chunk_index in range(num_chunks):
            start_byte = chunk_index * max_chunk_size
            end_byte = min(start_byte + max_chunk_size, file_size)
            chunk_data = audio_file_obj.read(end_byte - start_byte)

            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_wav_file:
              temp_wav_file_path = temp_wav_file.name  # Get the temporary file path

              # Write chunk data to the temporary WAV file
              with wave.open(temp_wav_file_path, "wb") as wav_file:
                  # Assuming the original audio format can be inferred (improve for unknown formats)
                  # Get audio parameters from the original file (if available)
                  original_wave_obj = wave.open(f"tmp/{audio_file}", "rb")
                  audio_params = (original_wave_obj.getnchannels(), original_wave_obj.getsampwidth(),
                                  original_wave_obj.getframerate())
                  original_wave_obj.close()  # Close the original file object

                  wav_file.setnchannels(audio_params[0])  # Number of channels
                  wav_file.setsampwidth(audio_params[1])  # Sample width (bytes per sample)
                  wav_file.setframerate(audio_params[2])  # Frame rate (samples per second)
                  wav_file.writeframes(chunk_data)

              # Use the temporary WAV file for transcription
              response = client.audio.transcriptions.create(
                  model="whisper-1",
                  file=open(temp_wav_file_path, "rb"),
                  language="en"  # Change to desired language code if needed
              )
              transcripts.append(response.text)

              # Delete the temporary WAV file after use
              os.remove(temp_wav_file_path)

    return "".join(transcripts)


def wrap_text(text, max_width, pdf):
    lines = []
    current_line = ""
    words = text.split()
    for word in words:
        if pdf.stringWidth(current_line + word, pdf._fontname, pdf._fontsize) < max_width:
            current_line += word + " "
        else:
            lines.append(current_line.strip())
            current_line = word + " "
    if current_line:
        lines.append(current_line.strip())
    return lines

def convertDocxToPdf(docx_path, pdf_path):
    document = Document(docx_path)
    pdf = canvas.Canvas(pdf_path, pagesize=letter)

    current_y_position = 792 - 72  # Starting position at the top of the page, considering bottom margin
    right_margin = 72  # Right margin in points
    bottom_margin = 72  # Bottom margin in points
    title_font_size = 14
    subtitle_font_size = 12
    normal_font_size = 12
    bullet_point = u'\u2022'  # Unicode character for bullet point

    for paragraph in document.paragraphs:
        style = paragraph.style.name

        if 'Heading' in style:
            # Handle titles and subtitles
            font_size = title_font_size if 'Heading 1' in style else subtitle_font_size
            pdf.setFont("Helvetica-Bold", font_size)
            text = paragraph.text
            lines = wrap_text(text, 912, pdf)
            for line in lines:
                pdf.drawString(72, current_y_position - pdf._leading - font_size, line)
                current_y_position -= pdf._leading + font_size
        else:
            # Handle bullet points
            if paragraph.style.name == 'List Bullet':
                pdf.setFont("Helvetica", normal_font_size)
                text = bullet_point + ' ' + paragraph.text
                lines = wrap_text(text, 72, pdf)
                for line in lines:
                    pdf.drawString(72, max(current_y_position - pdf._leading - normal_font_size, bottom_margin), line)
                    current_y_position -= pdf._leading + normal_font_size
            else:
                # Handle normal text
                pdf.setFont("Helvetica", normal_font_size)
                text = paragraph.text
                lines = wrap_text(text, 72, pdf)
                for line in lines:
                    pdf.drawString(72, max(current_y_position - pdf._leading - normal_font_size, bottom_margin), line)
                    current_y_position -= pdf._leading + normal_font_size

        # Check if we need to start a new page
        if current_y_position <= bottom_margin:
            pdf.showPage()
            current_y_position = 792 - bottom_margin  # Reset to the top of the new page, considering the bottom margin

    pdf.save()
